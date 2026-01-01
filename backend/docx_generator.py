"""
Word document generation utilities for medical notes and templates
"""
import io
import time
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_note_docx(template_name: str, note_data: Dict[str, Any]) -> io.BytesIO:
    """
    Generate Word document for a medical note
    
    Args:
        template_name: Name of the template used
        note_data: Dictionary containing note fields and values
        
    Returns:
        BytesIO: Word document as bytes
    """
    # Create Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Medical Note', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add template name
    if template_name:
        template_para = doc.add_paragraph(f'Template: {template_name}')
        template_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        template_run = template_para.runs[0]
        template_run.font.size = Pt(12)
        template_run.font.color.rgb = RGBColor(108, 117, 125)
        template_run.italic = True
    
    # Add horizontal line
    doc.add_paragraph('_' * 80)
    
    # Add note fields
    for key, value in note_data.items():
        # Skip internal fields
        if key.startswith('_'):
            continue
        
        # Format field name
        field_name = key.replace('_', ' ').title()
        
        # Field heading
        field_heading = doc.add_heading(field_name, level=2)
        field_heading.runs[0].font.color.rgb = RGBColor(102, 126, 234)
        
        # Field content
        if isinstance(value, dict):
            # Handle nested objects (recursively for triple-nested dicts)
            for sub_key, sub_value in value.items():
                sub_field_name = sub_key.replace('_', ' ').title()
                
                # Check if sub_value is also a dict (triple-nested)
                if isinstance(sub_value, dict):
                    # Add sub-heading
                    sub_heading = doc.add_paragraph(f"{sub_field_name}:")
                    sub_heading.runs[0].bold = True
                    sub_heading.runs[0].font.size = Pt(11)
                    
                    # Add nested items
                    for nested_key, nested_value in sub_value.items():
                        nested_field = nested_key.replace('_', ' ').title()
                        nested_para = doc.add_paragraph(f"  • {nested_field}: {nested_value}", style='List Bullet')
                        nested_para.runs[0].font.size = Pt(10)
                else:
                    # Regular nested field
                    content = doc.add_paragraph(f"{sub_field_name}: {sub_value}")
                    content.runs[0].font.size = Pt(11)
        else:
            content = doc.add_paragraph(str(value))
            content.runs[0].font.size = Pt(11)
        
        # Add spacing
        doc.add_paragraph()
    
    # Save to bytes
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io


def generate_template_docx(template_name: str, fields: Dict[str, Any]) -> io.BytesIO:
    """
    Generate Word document for a template structure
    
    Args:
        template_name: Name of the template
        fields: Dictionary containing template fields
        
    Returns:
        BytesIO: Word document as bytes
    """
    # Create Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading(template_name, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle
    subtitle = doc.add_paragraph('Medical Template Structure')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(108, 117, 125)
    subtitle_run.italic = True
    
    # Add horizontal line
    doc.add_paragraph('_' * 80)
    
    # Add fields
    for key, value in fields.items():
        # Field title
        field_heading = doc.add_heading(key.replace('_', ' ').title(), level=2)
        field_heading.runs[0].font.color.rgb = RGBColor(102, 126, 234)
        
        # Field description
        desc = doc.add_paragraph(value.get('description', 'No description provided'))
        desc_run = desc.runs[0]
        desc_run.font.size = Pt(11)
        
        # Field type
        type_para = doc.add_paragraph(f"Type: {value.get('type', 'text')}")
        type_run = type_para.runs[0]
        type_run.font.size = Pt(9)
        type_run.font.color.rgb = RGBColor(102, 126, 234)
        type_run.bold = True
        
        # Add spacing
        doc.add_paragraph()
    
    # Save to bytes
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io


def get_note_filename() -> str:
    """Generate filename for medical note"""
    return f"Medical_Note_{int(time.time())}.docx"


def get_template_filename(template_name: str) -> str:
    """Generate filename for template"""
    return f"{template_name.replace(' ', '_')}.docx"
