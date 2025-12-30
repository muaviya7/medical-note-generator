# Extract text from PDF and DOCX files
import os
import logging

logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_path):
    """
    Extract text from PDF using PyMuPDF
    
    Args:
        file_path: Path to PDF file
        
    Returns:
        str: Extracted text
    """
    try:
        import fitz  # PyMuPDF
        
        doc = fitz.open(file_path)
        text = ""
        
        for page in doc:
            text += page.get_text()
        
        doc.close()
        logger.info(f"Extracted {len(text)} chars from PDF")
        return text.strip()
        
    except Exception as e:
        logger.error(f"PDF extraction failed: {str(e)}")
        raise


def extract_text_from_docx(file_path):
    """
    Extract text from DOCX using python-docx
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        str: Extracted text
    """
    try:
        from docx import Document
        
        doc = Document(file_path)
        text = ""
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        logger.info(f"Extracted {len(text)} chars from DOCX")
        return text.strip()
        
    except Exception as e:
        logger.error(f"DOCX extraction failed: {str(e)}")
        raise


def extract_text(file_path):
    """
    Main function to extract text from file
    Detects file type and uses appropriate extractor
    
    Args:
        file_path: Path to file
        
    Returns:
        dict: {"success": bool, "text": str, "error": str}
    """
    result = {
        "success": False,
        "text": "",
        "error": ""
    }
    
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Check file size
        file_size = os.path.getsize(file_path)
        from .config import MAX_UPLOAD_SIZE
        
        if file_size > MAX_UPLOAD_SIZE:
            raise ValueError(f"File too large: {file_size} bytes (max {MAX_UPLOAD_SIZE})")
        
        # Detect file type
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            text = extract_text_from_pdf(file_path)
        elif ext == '.docx':
            text = extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
        
        if not text or len(text) < 10:
            raise ValueError("Extracted text is too short - document may be empty or scanned")
        
        result["success"] = True
        result["text"] = text
        logger.info(f"Successfully extracted text from {os.path.basename(file_path)}")
        
    except Exception as e:
        error_msg = str(e)
        logger.error(f"Text extraction error: {error_msg}")
        result["error"] = error_msg
    
    return result


# Test
