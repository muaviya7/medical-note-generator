from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.responses import JSONResponse, HTMLResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Dict, Any
import logging
import os
import json
import time
import shutil
import io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import modules and their models
from backend.text_cleaner import clean_text, CleanTextRequest, CleanTextResponse
from backend.note_generator import generate_note_from_text, load_template, GenerateNoteRequest, GenerateNoteResponse
from backend.template_generator import create_template_from_document, CreateTemplateResponse
from backend.transcription import transcribe_audio
from backend.note_formatter import format_medical_note, format_template_document
from backend.config import TEMPLATE_DIR, ALLOWED_EXTENSIONS, MAX_UPLOAD_SIZE

app = FastAPI(title="Medical Note Generator API", version="1.0.0")

# Add CORS middleware
ALLOWED_ORIGINS = os.getenv("ALLOWED_ORIGINS", "*").split(",")
app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,  # Lock down in production
    allow_credentials=True,
    allow_methods=["POST", "GET", "DELETE"],  # Only needed methods
    allow_headers=["*"],
)

# Mount static files
app.mount("/static", StaticFiles(directory="frontend/static"), name="static")
app.mount("/public", StaticFiles(directory="frontend/public"), name="public")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ============================================================================
# RESPONSE MODELS FOR COMBINED ENDPOINTS
# ============================================================================

class TranscribeAndCleanResponse(BaseModel):
    success: bool
    transcription: str = ""
    total_time: float = 0.0
    error: str | None = None

class DownloadTemplateRequest(BaseModel):
    template_name: str
    fields: Dict[str, Any]

class DownloadNoteRequest(BaseModel):
    template_name: str
    note_data: Dict[str, Any]

# ============================================================================
# FRONTEND & HEALTH CHECK
# ============================================================================

@app.get("/", response_class=HTMLResponse)
async def serve_frontend():
    """Serve the frontend HTML"""
    index_path = Path("frontend/public/index.html")
    if index_path.exists():
        return index_path.read_text(encoding="utf-8")
    return "<h1>Medical Note Generator</h1><p>Frontend not found</p>"

@app.get("/health")
def health_check():
    """API health check"""
    return {
        "status": "healthy",
        "service": "Medical Note Generator API",
        "version": "1.0.0"
    }

# ============================================================================
# API 1: TRANSCRIBE + CLEAN TEXT (Combined)
# ============================================================================

@app.post("/transcribe-and-clean", response_model=TranscribeAndCleanResponse, response_model_exclude_none=True)
async def transcribe_and_clean_audio(audio: UploadFile = File(...)):
    """
    Step 1: Upload audio file → Get transcription + cleaned text
    
    Args:
        audio: Audio file (mp3, wav, m4a, etc.)
    
    Returns:
        {
            "success": bool,
            "transcription": str,
            "total_time": float,
            "error": str (only on failure)
        }
    """
    try:
        total_start = time.time()
        
        # Save uploaded audio file temporarily
        temp_audio = f"temp_audio_{int(time.time())}{Path(audio.filename).suffix}"
        try:
            with open(temp_audio, "wb") as f:
                shutil.copyfileobj(audio.file, f)
            
            # Step 1 - Transcribe audio using Whisper
            transcription_start = time.time()
            transcription_result = transcribe_audio(temp_audio)
            transcription_time = time.time() - transcription_start
            
            if not transcription_result['success']:
                return TranscribeAndCleanResponse(
                    success=False,
                    error=f"Transcription failed: {transcription_result.get('error', 'Unknown error')}",
                    total_time=time.time() - total_start
                )
            
            transcription = transcription_result['text']
            logger.info(f"Transcription completed: {len(transcription)} chars in {transcription_time:.2f}s")
            
        finally:
            # Cleanup temp audio file
            if os.path.exists(temp_audio):
                try:
                    os.remove(temp_audio)
                except:
                    pass
        
        # Step 2 - Clean text
        cleaning_start = time.time()
        cleaned_result = clean_text(transcription)
        cleaning_time = time.time() - cleaning_start
        
        total_time = time.time() - total_start
        
        if cleaned_result['success']:
            return TranscribeAndCleanResponse(
                success=True,
                transcription=cleaned_result['formatted_text'],
                total_time=total_time
            )
        else:
            return TranscribeAndCleanResponse(
                success=False,
                error="Failed to clean transcribed text",
                total_time=total_time
            )
            
    except Exception as e:
        logger.error(f"Transcribe and clean error: {e}")
        return TranscribeAndCleanResponse(
            success=False,
            error=str(e),
            total_time=0.0
        )

# ============================================================================
# API 2: GENERATE MEDICAL NOTE
# ============================================================================

@app.post("/generate-note", response_model=GenerateNoteResponse, response_model_exclude_none=True)
async def generate_medical_note(request: GenerateNoteRequest):
    """
    Step 2: Generate structured medical note from cleaned text
    
    Args:
        request: {
            "cleaned_text": str,
            "template_name": str (template name from database)
        }
    
    Returns:
        {
            "success": bool,
            "medical_note": dict,
            "time_elapsed": float,
            "error": str (only on failure)
        }
    """
    try:
        start = time.time()
        
        # Generate note using template from database
        result = generate_note_from_text(request.cleaned_text, request.template_name)
        elapsed = time.time() - start
        
        if "error" in result:
            return GenerateNoteResponse(
                success=False,
                error=result["error"]
            )
        
        # Format the note as HTML
        formatted_html = format_medical_note(result, request.template_name)
        
        return GenerateNoteResponse(
            success=True,
            medical_note=result,
            time_elapsed=elapsed,
            formatted_html=formatted_html
        )
    except Exception as e:
        logger.error(f"Note generation error: {e}")
        return GenerateNoteResponse(
            success=False,
            error=str(e)
        )

# ============================================================================
# API 3: LIST TEMPLATES
# ============================================================================

@app.get("/templates")
async def list_all_templates():
    """
    Get list of all available templates
    
    Returns:
        {
            "success": bool,
            "templates": [{"name": str, "field_count": int, "created_at": str}]
        }
    """
    try:
        from backend.database import list_templates
        templates = list_templates()
        return {
            "success": True,
            "templates": templates
        }
    except Exception as e:
        logger.error(f"Error listing templates: {e}")
        return {
            "success": False,
            "error": str(e),
            "templates": []
        }

# ============================================================================
# API 4: CREATE TEMPLATE FROM DOCUMENT
# ============================================================================

@app.post("/create-template", response_model=CreateTemplateResponse, response_model_exclude_none=True)
async def create_template(document: UploadFile = File(...), template_name: str = Form(...)):
    """
    Step 3: Upload medical document (PDF/DOCX) → Generate template
    
    Args:
        document: PDF or DOCX medical form/document
        template_name: Name for the new template
    
    Returns:
        {
            "success": bool,
            "template_name": str,
            "fields": dict,
            "path": str,
            "time_elapsed": float,
            "error": str
        }
    """
    temp_file = None
    try:
        start = time.time()
        
        # Validate file extension
        file_ext = Path(document.filename).suffix.lower()
        if file_ext not in ALLOWED_EXTENSIONS:
            return CreateTemplateResponse(
                success=False,
                error=f"Invalid file type. Allowed: {', '.join(ALLOWED_EXTENSIONS)}"
            )
        
        # Save uploaded file temporarily
        temp_file = f"temp_{int(time.time())}{file_ext}"
        with open(temp_file, "wb") as f:
            shutil.copyfileobj(document.file, f)
        
        # Create template
        result = create_template_from_document(temp_file, template_name)
        elapsed = time.time() - start
        
        if result["success"]:
            # Format template as HTML
            formatted_html = format_template_document(result["fields"], result["template_name"])
            
            return CreateTemplateResponse(
                success=True,
                template_name=result["template_name"],
                fields=result["fields"],
                time_elapsed=elapsed,
                formatted_html=formatted_html,
                message=f"Template '{result['template_name']}' successfully added to database"
            )
        else:
            return CreateTemplateResponse(
                success=False,
                error=result["error"]
            )
    except Exception as e:
        logger.error(f"Template creation error: {e}")
        return CreateTemplateResponse(
            success=False,
            error=str(e)
        )
    finally:
        # Cleanup temp file
        if temp_file and os.path.exists(temp_file):
            try:
                os.remove(temp_file)
            except:
                pass

# ============================================================================
# API 5: DOWNLOAD TEMPLATE AS WORD
# ============================================================================

@app.post("/download-template")
async def download_template(request: DownloadTemplateRequest):
    """
    Download template as Word document
    
    Args:
        request: Template name and fields
    
    Returns:
        Word document file
    """
    try:
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading(request.template_name, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subtitle
        subtitle = doc.add_paragraph('Medical Template Structure')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.color.rgb = RGBColor(108, 117, 125)
        subtitle_run.italic = True
        
        # Add horizontal line
        doc.add_paragraph('_' * 80)
        
        # Add fields
        for key, value in request.fields.items():
            # Field title
            field_heading = doc.add_heading(key.replace('_', ' ').title(), level=2)
            field_heading.runs[0].font.color.rgb = RGBColor(102, 126, 234)
            
            # Field description
            desc = doc.add_paragraph(value.get('description', 'No description provided'))
            desc_run = desc.runs[0]
            desc_run.font.size = Pt(11)
            
            # Field type
            type_para = doc.add_paragraph(f"Type: {value.get('type', 'text')}")
            type_run = type_para.runs[0]
            type_run.font.size = Pt(9)
            type_run.font.color.rgb = RGBColor(102, 126, 234)
            type_run.bold = True
            
            # Add spacing
            doc.add_paragraph()
        
        # Save to bytes
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        # Return as downloadable file
        filename = f"{request.template_name.replace(' ', '_')}.docx"
        return StreamingResponse(
            doc_io,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except Exception as e:
        logger.error(f"Template download error: {e}")
        return JSONResponse(
            status_code=500,
            content={"error": str(e)}
        )

@app.post("/download-note")
async def download_note(request: DownloadNoteRequest):
    """
    Download medical note as Word document
    
    Args:
        request: Template name and note data
    
    Returns:
        Word document file
    """
    try:
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Medical Note', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add template name
        if request.template_name:
            template_para = doc.add_paragraph(f'Template: {request.template_name}')
            template_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            template_run = template_para.runs[0]
            template_run.font.size = Pt(12)
            template_run.font.color.rgb = RGBColor(108, 117, 125)
            template_run.italic = True
        
        # Add horizontal line
        doc.add_paragraph('_' * 80)
        
        # Add note fields
        for key, value in request.note_data.items():
            # Skip internal fields
            if key.startswith('_'):
                continue
            
            # Format field name
            field_name = key.replace('_', ' ').title()
            
            # Field heading
            field_heading = doc.add_heading(field_name, level=2)
            field_heading.runs[0].font.color.rgb = RGBColor(102, 126, 234)
            
            # Field content
            if isinstance(value, dict):
                # Handle nested objects
                for sub_key, sub_value in value.items():
                    sub_field_name = sub_key.replace('_', ' ').title()
                    content = doc.add_paragraph(f"{sub_field_name}: {sub_value}")
                    content.runs[0].font.size = Pt(11)
            else:
                content = doc.add_paragraph(str(value))
                content.runs[0].font.size = Pt(11)
            
            # Add spacing
            doc.add_paragraph()
        
        # Save to bytes
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        # Return as downloadable file
        filename = f"Medical_Note_{int(time.time())}.docx"
        return StreamingResponse(
            doc_io,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except Exception as e:
        logger.error(f"Note download error: {e}")
        return JSONResponse(
            status_code=500,
            content={"error": str(e)}
        )

# ============================================================================
# RUN SERVER
# ============================================================================

if __name__ == "__main__":
    import uvicorn
    from backend.config import RELOAD, APP_ENV
    
    # Get port and host from environment (Render provides these)
    port = int(os.getenv("PORT", 8000))
    host = os.getenv("HOST", "0.0.0.0")
    reload = os.getenv("RELOAD", "false").lower() == "true" if APP_ENV == "development" else False
    
    # Production-ready configuration
    uvicorn_config = {
        "host": host,
        "port": port,
        "reload": reload,
    }
    
    # Only include reload settings in development
    if APP_ENV == "development" and reload:
        uvicorn_config.update({
            "reload_dirs": ["./backend", "./"],
            "reload_excludes": [
                "*.pyc",
                "__pycache__",
                "*.log",
                "*.tmp",
                "temp_*",
                "*.db",
                "*.sqlite",
                ".git",
                "venv",
                ".venv",
                "models/*",
                "*.gguf"
            ]
        })
    
    logger.info(f"Starting server in {APP_ENV} mode on {host}:{port}")
    
    uvicorn.run(
        "main:app",
        **uvicorn_config
    )